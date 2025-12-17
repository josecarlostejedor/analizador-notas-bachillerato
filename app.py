import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openai
import io
import pdfplumber
import docx
from datetime import datetime
import re

# --- CONFIGURACI√ìN DE P√ÅGINA ---
st.set_page_config(
    page_title="Acta de Evaluaci√≥n", 
    page_icon="üéì", 
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- ESTILOS CSS ---
st.markdown("""
<style>
    .stTabs [data-baseweb="tab-list"] { gap: 2px; }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        white-space: pre-wrap;
        background-color: #f0f2f6;
        border-radius: 4px 4px 0px 0px;
        padding-top: 10px;
        padding-bottom: 10px;
    }
    .stTabs [aria-selected="true"] {
        background-color: #ffffff;
        border-top: 2px solid #4e8cff;
    }
    div[data-testid="stDataEditor"] {
        border: 2px solid #4e8cff;
        border-radius: 5px;
    }
</style>
""", unsafe_allow_html=True)

# --- GESTI√ìN DE ESTADO ---
if 'uploader_key' not in st.session_state:
    st.session_state.uploader_key = 0
if 'data' not in st.session_state:
    st.session_state.data = None

def reiniciar_app():
    st.session_state.data = None
    st.session_state.uploader_key += 1
    st.rerun()

# --- FUNCIONES DE EXTRACCI√ìN ---
def get_pdf_text_content(file):
    """Extrae el texto crudo manteniendo la disposici√≥n visual."""
    text_content = ""
    try:
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                # x_tolerance ayuda a mantener columnas separadas visualmente
                text_content += page.extract_text(x_tolerance=2, y_tolerance=2) + "\n"
        return text_content
    except Exception as e:
        return ""

def extract_text_from_docx(file):
    try:
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    except: return ""

# --- LIMPIEZA DE NOMBRES ---
def limpiar_nombre_alumno(texto):
    """
    1. Quita n√∫meros de lista al principio (1, 2...).
    2. Gira "APELLIDOS, NOMBRE" a "NOMBRE APELLIDOS".
    """
    if not isinstance(texto, str): return str(texto)
    
    # Limpiar espacios extra
    texto = texto.strip()
    
    # 1. Quitar n√∫mero inicial (ej: "1 ANTHONY" -> "ANTHONY")
    # Busca d√≠gitos al inicio seguidos de espacio, punto o guion
    texto = re.sub(r'^\d+[\.\-\s]+', '', texto)
    
    # 2. Reordenar nombre si hay coma
    if ',' in texto:
        partes = texto.split(',')
        if len(partes) >= 2:
            apellidos = partes[0].strip()
            nombre = partes[1].strip()
            return f"{nombre} {apellidos}"
            
    return texto

def process_data_with_ai(text_data, api_key, filename):
    if not text_data or len(text_data) < 10: 
        return None
        
    client = openai.OpenAI(api_key=api_key)
    
    # PROMPT DISE√ëADO PARA EVITAR EL ERROR DE COLUMNAS
    prompt = f"""
    Analiza el texto de este acta de evaluaci√≥n ('{filename}').
    
    ESTRUCTURA DEL PDF:
    1. Lista de alumnos (ej: "1. APELLIDO, NOMBRE"). 
       IMPORTANTE: El n√∫mero "1" es un √≠ndice, NO ES LA NOTA.
    2. Las notas (0-10) suelen estar en un bloque separado o al final de la l√≠nea.
    3. Asocia la fila de notas correcta a cada alumno.
    
    TAREA OBLIGATORIA:
    Genera datos separados por la barra vertical '|'. NO USES COMAS para separar columnas.
    Formato: Alumno|Materia|Nota
    
    REGLAS:
    - Alumno: Nombre COMPLETO tal cual aparece (ej: "PEREZ, JUAN"). NO resumas.
    - Materia: Abreviatura (ING1, EF, etc).
    - Nota: N√∫mero decimal.
    - Usa '|' como separador.
    
    Texto:
    {text_data[:20000]}
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}], 
            temperature=0
        )
        # Limpiar respuesta
        raw_content = response.choices[0].message.content
        csv_str = raw_content.replace("```csv", "").replace("```", "").strip()
        
        # Leemos usando el separador '|' para evitar errores con las comas del nombre
        df = pd.read_csv(io.StringIO(csv_str), sep='|', names=['Alumno', 'Materia', 'Nota'], engine='python')
        
        # Limpieza de nombres
        if 'Alumno' in df.columns:
            df['Alumno'] = df['Alumno'].apply(limpiar_nombre_alumno)
            
        return df
        
    except Exception as e:
        st.error(f"Error t√©cnico procesando datos: {e}")
        return None

# --- GENERACI√ìN DE TEXTOS AUTOM√ÅTICOS ---
def generar_comentario_individual(alumno, datos_alumno):
    suspensos = datos_alumno[datos_alumno['Nota'] < 5]
    num = len(suspensos)
    lista = suspensos['Materia'].tolist()
    txt = f"El alumno/a {alumno} tiene actualmente {num} materias suspensas."
    if num == 0: txt = "No tiene ninguna materia suspensa. ¬°Excelente trabajo! Se recomienda mantener la constancia."
    elif num == 1: txt += f" La materia es: {', '.join(lista)}. Recuperaci√≥n factible con refuerzo."
    elif num == 2: txt += f" Las materias son: {', '.join(lista)}. Situaci√≥n l√≠mite. Se aconseja organizaci√≥n urgente."
    else: txt += f" Las materias son: {', '.join(lista)}. Situaci√≥n preocupante que compromete la promoci√≥n."
    return txt

def generar_valoracion_detallada(res):
    txt = f"Nota media global: {res['media_grupo']:.2f}. "
    if res['pct_pasan'] >= 85: txt += "Promoci√≥n excelente."
    elif res['pct_pasan'] >= 70: txt += "Promoci√≥n satisfactoria."
    else: txt += "Promoci√≥n baja, requiere intervenci√≥n."
    return txt

# --- WORD INDIVIDUAL ---
def add_alumno_to_doc(doc, alumno, datos_alumno, media, suspensos, stats_mat):
    doc.add_heading(f'Informe Individual: {alumno}', 0)
    doc.add_paragraph(f"Nota Media: {media:.2f} | Materias Suspensas: {suspensos}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('An√°lisis y Recomendaciones', level=2)
    p = doc.add_paragraph(generar_comentario_individual(alumno, datos_alumno))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('Detalle de Calificaciones', level=2)
    t = doc.add_table(rows=1, cols=4); t.style = 'Table Grid'
    hdr = t.rows[0].cells
    hdr[0].text='Materia'; hdr[1].text='Nota'; hdr[2].text='Media Clase'; hdr[3].text='Dif.'
    
    medias = stats_mat.set_index('Materia')['Media'].to_dict()
    for _, row in datos_alumno.iterrows():
        c = t.add_row().cells
        c[0].text = str(row['Materia']); c[1].text = str(row['Nota'])
        mc = medias.get(row['Materia'], 0); c[2].text = f"{mc:.2f}"
        dif = row['Nota'] - mc; c[3].text = f"{dif:+.2f}"
        if row['Nota'] < 5:
            run = c[1].paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(255,0,0); run.bold = True

    # PIE DE P√ÅGINA
    doc.add_paragraph("\n\n")
    now = datetime.now()
    meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
    fecha_str = f"En Salamanca, a {now.day} de {meses[now.month-1]} de {now.year}"
    
    p_f = doc.add_paragraph(fecha_str); p_f.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("\n")
    p_s = doc.add_paragraph("El Tutor del grupo:"); p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_s.add_run("\n\n\n")
    p_s.add_run("D. Jos√© Carlos Tejedor Lorenzo").bold = True

def crear_informe_individual(alumno, datos_alumno, media, suspensos, stats_mat):
    doc = Document()
    add_alumno_to_doc(doc, alumno, datos_alumno, media, suspensos, stats_mat)
    bio = io.BytesIO(); doc.save(bio); bio.seek(0)
    return bio

def generar_informe_todos_alumnos(df, stats_al, stats_mat):
    doc = Document()
    for i, al in enumerate(stats_al['Alumno'].unique()):
        d_al = df[df['Alumno'] == al]
        info = stats_al[stats_al['Alumno'] == al].iloc[0]
        add_alumno_to_doc(doc, al, d_al, info['Media'], info['Suspensos'], stats_mat)
        if i < len(stats_al)-1: doc.add_page_break()
    bio = io.BytesIO(); doc.save(bio); bio.seek(0)
    return bio

# --- WORD GLOBAL Y PADRES ---
def generate_global_report(datos_resumen, plots, ranking_materias, centro, grupo):
    doc = Document()
    s = doc.sections[0]; s.orientation = WD_ORIENT.LANDSCAPE; s.page_width, s.page_height = s.page_height, s.page_width
    doc.add_heading(f'Informe de Evaluaci√≥n - {centro}', 0)
    doc.add_heading('Datos Generales', 1)
    doc.add_paragraph(f"Media: {datos_resumen['media_grupo']:.2f} | Promoci√≥n: {datos_resumen['pasan']} ({datos_resumen['pct_pasan']:.1f}%)")
    doc.add_paragraph(datos_resumen['valoracion']).italic = True
    doc.add_heading('Gr√°ficas', 1)
    if len(plots) >= 4:
        t = doc.add_table(rows=2, cols=2); t.autofit = True
        t.rows[0].cells[0].paragraphs[0].add_run().add_picture(plots[0], width=Inches(4.5))
        t.rows[0].cells[1].paragraphs[0].add_run().add_picture(plots[3], width=Inches(4.5))
        t.rows[1].cells[0].paragraphs[0].add_run().add_picture(plots[2], width=Inches(4.5))
        t.rows[1].cells[1].paragraphs[0].add_run().add_picture(plots[1], width=Inches(4.5))
    bio = io.BytesIO(); doc.save(bio); bio.seek(0)
    return bio

def generate_parents_report(res, stats_mat, plot_suspensos, plot_pct_materias):
    doc = Document()
    s = doc.sections[0]; s.orientation = WD_ORIENT.LANDSCAPE; s.page_width, s.page_height = s.page_height, s.page_width
    doc.add_heading('RESUMEN DE EVALUACI√ìN PARA FAMILIAS', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = doc.add_table(rows=1, cols=2); t.autofit = False
    t.columns[0].width = Inches(5); t.columns[1].width = Inches(5)
    
    c1 = t.rows[0].cells[0].paragraphs[0]
    c1.add_run("Resumen estad√≠stico.\n\n").italic = True
    c1.add_run(f"‚Ä¢ Promocionan: {res['pasan']} ({res['pct_pasan']:.1f}%)\n")
    c1.add_run(f"‚Ä¢ No promocionan: {res['no_pasan']} ({res['pct_no_pasan']:.1f}%)\n")
    c1.add_run(f"‚Ä¢ Media suspensos: {res['media_suspensos_grupo']:.2f}\n\n")
    c1.add_run("Aprobados por materia:\n").bold = True
    for _, row in stats_mat.iterrows(): c1.add_run(f"- {row['Materia']}: {row['Pct_Aprobados']:.1f}%\n")
    
    c2 = t.rows[0].cells[1]
    c2.paragraphs[0].add_run("Materias no superadas:\n").bold = True
    c2.paragraphs[0].add_run().add_picture(plot_suspensos, width=Inches(4.5))
    c2.add_paragraph("\n% Suspensos por Materia:\n").bold = True
    c2.paragraphs[1].add_run().add_picture(plot_pct_materias, width=Inches(4.5))
    
    bio = io.BytesIO(); doc.save(bio); bio.seek(0)
    return bio

# --- INTERFAZ ---
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/2991/2991148.png", width=50)
    st.title("Configuraci√≥n")
    api_key = st.text_input("üîë API Key OpenAI", type="password")
    st.markdown("---")
    centro = st.text_input("Centro", "IES Luc√≠a de Medrano")
    grupo = st.text_input("Grupo", "1¬∫ BACH 7")
    curso = st.text_input("Curso", "2024-2025")
    st.markdown("---")
    uploaded_files = st.file_uploader("üìÇ Subir Actas", type=['xlsx', 'pdf', 'docx', 'doc'], accept_multiple_files=True, key=f"up_{st.session_state.uploader_key}")
    
    if uploaded_files and st.session_state.data is None:
        if st.button("Analizar Datos", type="primary"):
            if not api_key: st.error("Falta API Key")
            else:
                dfs = []
                bar = st.progress(0)
                for i, f in enumerate(uploaded_files):
                    df_t = None
                    if f.name.endswith('.xlsx'):
                        try: df_t = pd.read_excel(f)
                        except: pass
                    elif f.name.endswith('.pdf'):
                        # Extracci√≥n robusta de texto
                        txt = get_pdf_text_content(f)
                        if txt: df_t = process_data_with_ai(txt, api_key, f.name)
                    elif 'doc' in f.name:
                        txt = extract_text_from_docx(f)
                        if txt: df_t = process_data_with_ai(txt, api_key, f.name)
                    
                    if df_t is not None: dfs.append(df_t)
                    bar.progress((i+1)/len(uploaded_files))
                
                if dfs:
                    st.session_state.data = pd.concat(dfs, ignore_index=True)
                    st.rerun()
                else: st.error("No se extrajeron datos. Revisa el archivo.")

    if st.session_state.data is not None:
        if st.button("üîÑ Subir nuevo"): reiniciar_app()

st.title("Acta de Evaluaci√≥n")
col_b1, col_b2, col_b3 = st.columns([1,1,1])
col_b1.info(f"üè´ **Centro:** {centro}")
col_b2.info(f"üë• **Grupo:** {grupo}")
col_b3.info(f"üìÖ **Curso:** {curso}")

if st.session_state.data is not None:
    df = st.session_state.data
    # Normalizaci√≥n de columnas autom√°tica (si ya vienen bien del CSV)
    if len(df.columns) >= 3:
        df.columns = ['Alumno', 'Materia', 'Nota']
    
    # Limpieza y c√°lculos
    df = df.drop_duplicates(subset=['Alumno', 'Materia'], keep='last')
    df['Nota'] = pd.to_numeric(df['Nota'], errors='coerce')
    df['Aprobado'] = df['Nota'] >= 5
    
    stats_al = df.groupby('Alumno').agg(Suspensos=('Nota', lambda x: (x<5).sum()), Media=('Nota', 'mean')).reset_index()
    stats_mat = df.groupby('Materia').agg(Total=('Nota', 'count'), Aprobados=('Aprobado', 'sum'), Suspensos=('Nota', lambda x: (x<5).sum()), Media=('Nota', 'mean')).reset_index()
    stats_mat['Pct_Aprobados'] = (stats_mat['Aprobados']/stats_mat['Total'])*100
    stats_mat['Pct_Suspensos'] = (stats_mat['Suspensos']/stats_mat['Total'])*100
    
    total = len(stats_al); media_gr = df['Nota'].mean()
    cero = stats_al[stats_al['Suspensos']==0].shape[0]
    uno = stats_al[stats_al['Suspensos']==1].shape[0]
    dos = stats_al[stats_al['Suspensos']==2].shape[0]
    tres = stats_al[stats_al['Suspensos']==3].shape[0]
    mas_tres = stats_al[stats_al['Suspensos']>3].shape[0]
    pasan = cero+uno+dos; no_pasan = tres+mas_tres
    base = total if total>0 else 1
    
    res = {'total_alumnos': total, 'media_grupo': media_gr, 'media_suspensos_grupo': stats_al['Suspensos'].mean(),
           'pasan': pasan, 'pct_pasan': (pasan/base)*100, 'no_pasan': no_pasan, 'pct_no_pasan': (no_pasan/base)*100, 'pct_mas_dos': ((tres+mas_tres)/base)*100}
    res['valoracion'] = generar_valoracion_detallada(res)

    tab1, tab2, tab3, tab4, tab5 = st.tabs(["üìä General", "üìö Materias", "üéì Editor", "üìÑ Informes", "üë®‚Äçüë©‚Äçüëß Padres"])
    
    with tab1:
        st.metric("Media Grupo", f"{media_gr:.2f}")
        c1,c2 = st.columns(2)
        with c1: 
            fig_p, ax_p = plt.subplots(figsize=(4,3)); ax_p.pie([pasan, no_pasan], labels=['S√≠', 'No'], autopct='%1.f%%', colors=['#2ecc71','#e74c3c']); st.pyplot(fig_p)
        with c2:
            fig_b, ax_b = plt.subplots(figsize=(4,3)); ax_b.bar(['0','1','2','3','>3'], [cero,uno,dos,tres,mas_tres], color='#3498db'); st.pyplot(fig_b)
            
        fig_d, ax_d = plt.subplots(figsize=(5,4)); bars_d = ax_d.bar(['0', '1', '2', '>2'], [cero, uno, dos, tres+mas_tres], color='#3498db'); ax_d.bar_label(bars_d)
        fig_m, ax_m = plt.subplots(figsize=(10,5)); d_gf = stats_mat.sort_values('Media', ascending=False); bars_m = ax_m.bar(d_gf['Materia'], d_gf['Media'], color='#9b59b6'); ax_m.bar_label(bars_m, fmt='%.2f')
        fig_pr, ax_pr = plt.subplots(figsize=(8,3)); ax_pr.bar(['S√≠', 'No'], [pasan, no_pasan], color=['green', 'red'])
        
        plots = []
        for f in [fig_p, fig_d, fig_m, fig_pr]:
            buf = io.BytesIO(); f.savefig(buf, format='png', bbox_inches='tight'); buf.seek(0); plots.append(buf)
        
        if st.button("üìÑ Informe General Word"):
            st.download_button("Descargar", generate_global_report(res, plots, stats_mat, centro, grupo), f"Global_{grupo}.docx", type="primary")

    with tab2: st.dataframe(stats_mat.style.format({'Pct_Aprobados':'{:.1f}%'}), use_container_width=True)
    
    with tab3:
        st.markdown("### Editor"); piv = df.pivot_table(index='Alumno', columns='Materia', values='Nota', aggfunc='first')
        ed = st.data_editor(piv, use_container_width=True)
        if st.button("üîÑ Recalcular"):
            try:
                new_df = ed.reset_index().melt(id_vars='Alumno', var_name='Materia', value_name='Nota')
                new_df.dropna(subset=['Nota'], inplace=True); st.session_state.data = new_df; st.rerun()
            except: pass

    with tab4:
        c1, c2 = st.columns(2)
        with c1:
            sel = st.selectbox("Alumno", stats_al['Alumno'].unique())
            if sel:
                inf = stats_al[stats_al['Alumno']==sel].iloc[0]
                st.info(generar_comentario_individual(sel, df[df['Alumno']==sel]))
                st.download_button("Descargar", crear_informe_individual(sel, df[df['Alumno']==sel], inf['Media'], inf['Suspensos'], stats_mat), f"{sel}.docx")
        with c2:
            if st.button("üöÄ Informe TODOS"):
                st.download_button("Descargar ZIP", generar_informe_todos_alumnos(df, stats_al, stats_mat), f"Todos_{grupo}.docx", type="primary")

    with tab5:
        fig_p1, ax_p1 = plt.subplots(figsize=(6,4))
        bars_p = ax_p1.bar(['0','1','2','3','>3'], [cero,uno,dos,tres,mas_tres], color=['#2ecc71','#f1c40f','#e67e22','#e74c3c','#c0392b'])
        ax_p1.bar_label(bars_p)
        
        fig_p2, ax_p2 = plt.subplots(figsize=(6,4))
        df_p2 = stats_mat.sort_values('Pct_Suspensos')
        ax_p2.barh(df_p2['Materia'], df_p2['Pct_Suspensos'], color='#3498db')
        
        c1,c2 = st.columns(2); c1.pyplot(fig_p1); c2.pyplot(fig_p2)
        b1 = io.BytesIO(); fig_p1.savefig(b1, format='png', bbox_inches='tight'); b1.seek(0)
        b2 = io.BytesIO(); fig_p2.savefig(b2, format='png', bbox_inches='tight'); b2.seek(0)
        if st.button("üìÑ Word Padres"):
            st.download_button("Descargar", generate_parents_report(res, stats_mat, b1, b2), f"Padres_{grupo}.docx", type="primary")
else: st.info("Sube archivo")
